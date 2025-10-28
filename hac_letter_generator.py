import streamlit as st
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import _Cell, Table
from docx.text.run import Run
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import pandas as pd
from io import BytesIO
import os
import re
import zipfile


st.set_page_config(page_title="HAC Letter Generator", layout="centered")
st.title("📄 Hybrid Asset Custody Letter Generator")

# -----------------------------
# Template registry
# -----------------------------
template_options = {
    "Letter of Appointment": "templates/Letter of Appointment.docx",
    "Letter of Instruction": "templates/Letter of Instruction.docx",
    "Letter of Set-Off": "templates/Letter of Set-Off.docx",
    "Letter of Affirmation (Quarterly)": "templates/Letter of Affirmation - Quarterly.docx",
    "Letter of Affirmation (Yearly)": "templates/Letter of Affirmation - Yearly.docx",
    "Letter of Welcome": "templates/Letter of Welcome.docx",
    "Letter of Agreement": "templates/Letter of Agreement.docx",
    "Letter of Acknowledgement": "templates/Letter of Acknowledgement.docx",
    "Letter of Dividend": "templates/Letter of Dividend.docx",
    "Letter of Lucky Draw": "templates/Letter of Lucky Draw.docx",
    "Trust Deed": "templates/Trust Deed.docx",
    "Commission Letter (7th)": "templates/Commission Letter (7th).docx",
    "HAC AGREEMENT": "templates/HAC Agreement.docx",
    "Letter of Dividend Payout": "templates/Letter of Dividend Payout.docx",
    "Letter of Dividend Payout (RPS)": "templates/Letter of Dividend Payout (RPS).docx",
    "Letter of Dividend Payout (GGFT)": "templates/Letter of Dividend Payout (GGFT).docx",
    "Mutual NDA": "templates/Mutual NDA.docx",
    "Marketing Consultancy and Services Agreement": "templates/Marketing Consultancy and Services Agreement.-v2.0-IND.docx",
}

# -----------------------------
# Utilities
# -----------------------------

def norm(s: str) -> str:
    """Trim and collapse inner spaces into one space."""
    return re.sub(r"\s+", " ", s.strip()) if isinstance(s, str) else ""

def norm_multiline(s: str) -> str:
    """Preserve intended line breaks from Excel cells."""
    if s is None:
        return ""
    s = str(s)
    s = s.replace("\r\n", "\n").replace("\r", "\n").replace("\\n", "\n")
    lines = [ln.strip() for ln in s.split("\n")]
    return "\n".join(lines).strip()

def auto_multiline_address(s: str) -> str:
    """
    If address is a single line, split on commas into neat lines.
    Keep existing line breaks if already present. Preserve commas at line ends.
    """
    if s is None:
        return ""
    s = str(s)
    if "\n" in s:
        s = s.replace("\r\n", "\n").replace("\r", "\n")
        return "\n".join(part.strip() for part in s.split("\n") if part.strip())
    parts = [p.strip() for p in s.split(",") if p.strip()]
    return "\n".join(f"{p}," if i < len(parts) - 1 else p for i, p in enumerate(parts))

def extract_placeholders(doc: Document) -> set:
    """Find all {{PLACEHOLDER}} in paragraphs and tables."""
    placeholder_pattern = re.compile(r"{{(.*?)}}")
    found = set()
    for para in doc.paragraphs:
        found.update(placeholder_pattern.findall(para.text))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    found.update(placeholder_pattern.findall(para.text))
    return found

def extract_placeholders_from_docx_bytes(docx_bytes: bytes) -> set:
    """Scan all word/*.xml parts for placeholders, including headers/footers/shapes."""
    pattern = re.compile(r"{{(.*?)}}")
    found = set()
    with BytesIO(docx_bytes) as in_mem:
        with zipfile.ZipFile(in_mem, 'r') as zin:
            for item in zin.infolist():
                if item.filename.startswith('word/') and item.filename.endswith('.xml'):
                    try:
                        xml_text = zin.read(item.filename).decode('utf-8')
                    except UnicodeDecodeError:
                        continue
                    found.update(pattern.findall(xml_text))
    return found

def replace_in_docx_bytes(docx_bytes: bytes, mapping: dict, placeholders_upper_map: dict) -> bytes:
    """Low-level XML replace across word/*.xml parts to catch text in shapes, headers, and footers."""
    in_mem = BytesIO(docx_bytes)
    out_mem = BytesIO()
    with zipfile.ZipFile(in_mem, 'r') as zin, zipfile.ZipFile(out_mem, 'w', zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data_bytes = zin.read(item.filename)
            if item.filename.startswith('word/') and item.filename.endswith('.xml'):
                try:
                    xml_text = data_bytes.decode('utf-8')
                except UnicodeDecodeError:
                    zout.writestr(item, data_bytes)
                    continue
                for orig, up in placeholders_upper_map.items():
                    ph = '{{' + orig + '}}'
                    val = mapping.get(up, '') or ''
                    if ph in xml_text:
                        # also soften tabs adjacent to placeholders in these parts
                        xml_text = xml_text.replace("\t" + ph, ph).replace(ph + "\t", ph)
                        xml_text = xml_text.replace(ph, str(val))
                data_bytes = xml_text.encode('utf-8')
            zout.writestr(item, data_bytes)
    out_mem.seek(0)
    return out_mem.getvalue()

# ---------- Formatting helpers ----------

def _copy_run_format(src: Run, dst: Run) -> None:
    """Copy key inline styles from src to dst."""
    dst.bold = src.bold
    dst.italic = src.italic
    dst.underline = src.underline
    if src.font is not None:
        if src.font.size:
            dst.font.size = src.font.size
        if src.font.name:
            dst.font.name = src.font.name
        try:
            rgb = src.font.color.rgb
            if isinstance(rgb, RGBColor):
                dst.font.color.rgb = rgb
        except Exception:
            pass

def _add_run_copy_style(paragraph: Paragraph, src_run: Run, text: str = "", add_line_break: bool = False) -> Run:
    """Add a run, cloning inline style from src_run, and optionally add a line break first."""
    r = paragraph.add_run()
    r.bold = src_run.bold
    r.italic = src_run.italic
    r.underline = src_run.underline
    if src_run.font is not None:
        if src_run.font.size:
            r.font.size = src_run.font.size
        if src_run.font.name:
            r.font.name = src_run.font.name
        try:
            if src_run.font.color and src_run.font.color.rgb:
                r.font.color.rgb = src_run.font.color.rgb
        except Exception:
            pass
    if add_line_break:
        r.add_break()
    if text:
        r.add_text(text)
    return r

def _walk_block_items(doc: Document):
    """Yield all paragraphs from the document, including those in tables."""
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p

def _to_2dp(value: str) -> str:
    """Sanitize a numeric-like string and format to two decimals."""
    s = str(value).strip()
    s = s.replace(",", "")
    s = re.sub(r"[^0-9.\-]", "", s)
    if s in {"", ".", "-", "-."}:
        return ""
    try:
        return f"{float(s):.2f}"
    except ValueError:
        return ""

# -------- Address paragraph helpers (tab cleanup) --------

def _clear_tab_stops(paragraph: Paragraph) -> None:
    """Remove all <w:tabs> descendants from this paragraph, no XPath needed."""
    tabs_tag = qn('w:tabs')
    for el in list(paragraph._element.iter()):
        if el.tag == tabs_tag:
            parent = el.getparent()
            if parent is not None:
                parent.remove(el)

def _strip_tabs(s: str) -> str:
    return s.replace("\t", " ")

# -------- Multiline paragraph rebuild (handles ADDRESS neatly) --------

def _replace_multiline_paragraph(paragraph: Paragraph, placeholders_upper: dict, data_map: dict) -> bool:
    """
    For values containing newline(s), rebuild the paragraph cleanly:
    - clear tab stops and tab characters
    - insert hard line breaks
    - keep formatting by cloning from a source run
    """
    text = paragraph.text
    did = False

    src_run = paragraph.runs[0] if paragraph.runs else paragraph.add_run("")

    for orig, up in placeholders_upper.items():
        ph = f"{{{{{orig}}}}}"
        value = str(data_map.get(up, ""))
        if ph in text and ("\n" in value):
            did = True

            # Clean tabs and tab stops to prevent column effects
            text = _strip_tabs(text.replace("\t" + ph, ph).replace(ph + "\t", ph))
            _clear_tab_stops(paragraph)
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.page_break_before = False

            before, after = text.split(ph, 1)

            # Clear paragraph
            for r in list(paragraph.runs):
                r.text = ""
            paragraph.text = ""

            # Write "before" without forcing a newline
            before = before.rstrip()
            if before:
                _add_run_copy_style(paragraph, src_run, before)

            # Write multiline value
            lines = [ln.rstrip() for ln in value.split("\n")]
            if lines:
                _add_run_copy_style(paragraph, src_run, lines[0])
                for line in lines[1:]:
                    _add_run_copy_style(paragraph, src_run, line, add_line_break=True)

            # Write "after" on a new visual line only if text exists
            after = after.lstrip()
            if after:
                _add_run_copy_style(paragraph, src_run, after, add_line_break=True)

            break
    return did

# -------- Cross-run replacer for single-line placeholders --------

def _replace_placeholders_across_runs(paragraph: Paragraph, placeholders_upper: dict, data_map: dict) -> None:
    """
    Replace {{PLACEHOLDER}} even when split across runs.
    Handles single-line values. Preserves formatting of the starting run.
    """
    if not paragraph.runs:
        return

    runs = paragraph.runs
    concat = ""
    spans = []
    for idx, r in enumerate(runs):
        start = len(concat)
        concat += r.text
        spans.append((idx, start, len(concat)))

    def _find_span(pos: int):
        for idx, s, e in spans:
            if s <= pos < e:
                return idx, s, e
        return spans[-1][0], spans[-1][1], spans[-1][2]

    changed = True
    while changed:
        changed = False
        for orig, up in placeholders_upper.items():
            ph = "{{" + orig + "}}"
            pos = concat.find(ph)
            if pos == -1:
                continue

            value = str(data_map.get(up, ""))
            if "\n" in value:
                # multiline handled by the paragraph rebuild pass
                continue

            changed = True
            start_pos = pos
            end_pos = pos + len(ph)

            rs_idx, rs_s, _ = _find_span(start_pos)
            re_idx, re_s, _ = _find_span(end_pos - 1)

            before = runs[rs_idx].text[: start_pos - rs_s]
            after = runs[re_idx].text[end_pos - re_s :]

            runs[rs_idx].text = before + value
            for i in range(rs_idx + 1, re_idx):
                runs[i].text = ""
            if re_idx != rs_idx:
                runs[re_idx].text = after

            # rebuild concat
            concat = ""
            spans = []
            for idx, r in enumerate(paragraph.runs):
                s = len(concat)
                concat += r.text
                spans.append((idx, s, len(concat)))
            break

# -------- Document cleanup to avoid blank last page --------

def _strip_page_breaks(doc: Document) -> None:
    """Remove manual page breaks and 'page break before' flags."""
    # Turn off "page break before" everywhere
    for p in doc.paragraphs:
        if p.paragraph_format:
            p.paragraph_format.page_break_before = False

    # Remove manual page-break BRs inside runs
    br_tag = qn('w:br')
    type_attr = qn('w:type')
    for p in doc.paragraphs:
        for r in p.runs:
            for el in list(r._r):
                if el.tag == br_tag and el.get(type_attr) == 'page':
                    r._r.remove(el)

def _remove_trailing_blank_paragraphs(doc: Document) -> None:
    """Drop empty tail paragraphs that can trigger an extra page."""
    while doc.paragraphs:
        p = doc.paragraphs[-1]
        txt = p.text.strip()
        only_empty_runs = all((run.text.strip() == "") for run in p.runs)
        if txt == "" and only_empty_runs:
            p._element.getparent().remove(p._element)
        else:
            if p.paragraph_format:
                p.paragraph_format.page_break_before = False
            break

# -----------------------------
# UI — Excel upload
# -----------------------------
uploaded_excel = st.file_uploader("📈 Upload Master Excel File (All Letters)", type=["xlsx"], key="excel_upload")

# Session state for generated docs
if "generated_docs" not in st.session_state:
    st.session_state["generated_docs"] = []

if uploaded_excel:
    st.success("✅ File uploaded. Click the button to generate all letters.")
    debug = st.checkbox("🔍 Debug mode (show placeholders & mapping)")

    if st.button("Generate Letters"):
        try:
            df = pd.read_excel(uploaded_excel)
        except Exception as e:
            st.error(f"Failed to read Excel: {e}")
            st.stop()

        # Normalize headers to UPPER
        df.columns = [col.strip().upper() for col in df.columns]

        st.session_state["generated_docs"] = []
        ok_count = 0

        for index, row in df.iterrows():
            try:
                # Build raw data from row (no normalization yet)
                data = {}
                for k, v in row.items():
                    key = str(k).strip().upper()
                    val = "" if pd.isna(v) else str(v)
                    data[key] = val

                # Peek the letter type early for conditional logic
                letter_type_raw = str(data.get("LETTER_TYPE", "")).strip()
                letter_type_upper = letter_type_raw.upper()
                is_dividend = "DIVIDEND" in letter_type_upper

                # Normalize values, with conditional address handling
                for k in list(data.keys()):
                    if "ADDRESS" in k:
                        if is_dividend:
                            data[k] = auto_multiline_address(data[k])
                        else:
                            data[k] = norm_multiline(data[k])
                    else:
                        data[k] = norm(data[k])

                # --- Amount formatting fix (2 decimals, sanitize inputs) ---
                for field in ["AMOUNT", "DIVIDEND", "ACCUMULATED", "TRUST_CAPITAL"]:
                    if field in data and data[field]:
                        data[field] = _to_2dp(data[field])
                # --- End of amount fix ---

                # --- Add tab spacing for MONTH placeholder ---
                if "MONTH" in data and re.match(r"^[A-Za-z]+\s+\d{4}$", data["MONTH"]):
                    # Converts "October 2025" → "October\t2025"
                    data["MONTH"] = re.sub(r"\s+(\d{4})", r"\t\1", data["MONTH"])
                # --- End of month tab fix ---

                # Final normalized types for downstream logic
                letter_type = norm(letter_type_raw)
                payout_type = norm(data.get("PAYOUT_TYPE", ""))

                if not letter_type:
                    st.warning(f"Row {index+1}: LETTER_TYPE is empty. Skipping.")
                    continue

                # Build template key safely
                template_key = f"{letter_type} ({payout_type})" if letter_type == "Letter of Affirmation" else letter_type
                template_path = template_options.get(template_key)

                if not template_path:
                    st.warning(f"Row {index+1}: Template key not found → '{template_key}'. Check spelling/casing/spaces.")
                    continue
                if not os.path.exists(template_path):
                    st.warning(f"Row {index+1}: Template path does not exist → {template_path}")
                    continue

                # Load Word template
                with open(template_path, "rb") as f:
                    template_bytes = f.read()
                doc = Document(BytesIO(template_bytes))

                # Detect placeholders in this template (original case)
                placeholders_orig = extract_placeholders(doc)
                placeholders_upper = {p: p.upper() for p in placeholders_orig}

                # Ensure every placeholder exists in data (default to empty string)
                for orig, up in placeholders_upper.items():
                    if up not in data:
                        data[up] = ""

                if debug:
                    xml_ph = extract_placeholders_from_docx_bytes(template_bytes)
                    sample = ", ".join(sorted(list(xml_ph))[:12])
                    if len(xml_ph) > 12:
                        sample += ", …"
                    st.info(
                        f"Row {index+1}: template='{template_key}'\n"
                        f"Path='{template_path}'\n"
                        f"Placeholders (docx XML)={len(xml_ph)} → {sample}"
                    )

                # -----------------------------
                # Pass A: Paragraph-level multiline replacer (ONLY for Dividend)
                # -----------------------------
                if is_dividend:
                    for para in _walk_block_items(doc):
                        if "{{" in para.text:
                            _replace_multiline_paragraph(para, placeholders_upper, data)

                # -----------------------------
                # Pass B: Cross-run replacement for single-line values (all)
                # -----------------------------
                for para in _walk_block_items(doc):
                    if "{{" in para.text:
                        _replace_placeholders_across_runs(para, placeholders_upper, data)

                # Clean document to avoid stray blank pages (ONLY for Dividend)
                if is_dividend:
                    _strip_page_breaks(doc)
                    _remove_trailing_blank_paragraphs(doc)

                # Save docx first
                inter_buffer = BytesIO()
                doc.save(inter_buffer)
                inter_bytes = inter_buffer.getvalue()

                # -----------------------------
                # Pass C: Low-level XML replacement for headers/footers/shapes (all)
                # -----------------------------
                final_bytes = replace_in_docx_bytes(inter_bytes, data, placeholders_upper)

                name_part = data.get("NAME", "Client")
                filename = f"{name_part} - {template_key}.docx"

                st.session_state["generated_docs"].append({
                    "filename": filename,
                    "buffer": final_bytes,
                })
                ok_count += 1

            except Exception as e:
                st.error(f"Row {index+1} failed: {e}")

        if ok_count == 0:
            st.warning("No files were generated. Check the warnings/errors above.")
        else:
            st.success(f"✅ Generated {ok_count} file(s). See download buttons below.")

# -----------------------------
# Downloads & housekeeping
# -----------------------------
if st.session_state.get("generated_docs"):
    def safe_filename(name: str) -> str:
        """Replace illegal path characters so ZIP doesn’t create folders."""
        return re.sub(r'[\\/*?:"<>|]', "_", name)

    # Build a ZIP in memory containing all generated files
    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        for file in st.session_state["generated_docs"]:
            safe_name = safe_filename(file["filename"])
            zf.writestr(safe_name, file["buffer"])
    zip_buffer.seek(0)

    # Download-all button
    st.download_button(
        label="📦 Download all as ZIP",
        data=zip_buffer,
        file_name="HAC_Letters.zip",
        mime="application/zip",
        key="download_all_zip",
    )

    # Individual file buttons
    for i, file in enumerate(st.session_state["generated_docs"]):
        st.download_button(
            label=f"🗕️ Download: {file['filename']}",
            data=file["buffer"],
            file_name=file["filename"],
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key=f"download_{i}",
        )

    if st.button("🗑️ Clear All"):
        st.session_state["generated_docs"] = []
        st.rerun()
else:
    st.warning("📌 Please upload the Excel file to begin.")
